"""Generate the Player's Guide Word document for the Stock Simulator.

Run from the repo root:
    python scripts/build_players_guide.py

Writes docs/Players_Guide.docx.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "Players_Guide.docx"

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x55, 0x55, 0x55)


def set_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def style_heading(paragraph, color: RGBColor = ACCENT) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = color


def add_title(doc: Document, text: str, subtitle: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(text)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(subtitle)
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = MUTED


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    style_heading(h)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            label, body = item
            run = p.add_run(label)
            run.bold = True
            p.add_run(body)
        else:
            p.add_run(item)


def add_numbered(doc: Document, items) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.color.rgb = ACCENT
    p.add_run(text)


def add_table(doc: Document, headers, rows) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def build() -> None:
    doc = Document()
    set_base_style(doc)

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    add_title(
        doc,
        "Stock Simulator Player's Guide",
        "How to trade, join competitions, and compete as a team",
    )

    add_paragraph(
        doc,
        "Welcome to the Stock Simulator! This guide walks you through everything you "
        "need to know to trade with virtual cash, track your portfolio, and compete "
        "against your classmates as an individual or on a team.",
    )

    # ----- Quick Start -----
    add_heading(doc, "Quick Start", level=1)
    add_numbered(
        doc,
        [
            "Register for an account with your username, email, and password.",
            "Log in. Your Global account is ready with $100,000 in virtual cash.",
            "Ask your teacher for the Competition Code to join their competition.",
            "Join the competition as an individual, or create/join a team first.",
            "Start trading. Watch the leaderboard and try to beat your classmates!",
        ],
    )

    # ----- Accounts -----
    add_heading(doc, "Your Three Account Types", level=1)
    add_paragraph(
        doc,
        "You can hold up to three kinds of accounts at the same time. Each account "
        "is completely separate: it has its own cash balance, holdings, trade "
        "history, and performance metrics. Switching accounts in the app changes "
        "which portfolio you are trading.",
    )

    add_table(
        doc,
        headers=["Account", "When You Get It", "Starting Cash", "Used For"],
        rows=[
            (
                "Global",
                "Automatically, on registration",
                "$100,000",
                "Free-play practice trading any time, independent of any class or competition.",
            ),
            (
                "Competition",
                "When you join a teacher's competition as an individual",
                "$100,000",
                "Individual performance in a specific competition with a set start/end date.",
            ),
            (
                "Team",
                "When your team joins a competition",
                "$100,000",
                "Shared team portfolio — all teammates trade out of the same pool of cash.",
            ),
        ],
    )

    add_callout(
        doc,
        "Important",
        "The three accounts never mix. A trade you place in your Global account has no "
        "effect on your Competition or Team balances, and vice versa. Always confirm which "
        "account is selected before placing an order.",
    )

    # ----- Registration -----
    add_heading(doc, "Registration and Sign-In", level=1)
    add_bullets(
        doc,
        [
            ("Register: ", "Create your account with a unique username, a valid email, and a password. Emails and usernames must be unique."),
            ("Log in: ", "Use the credentials you set at registration. Your dashboard loads all three account types you belong to."),
            ("Forgot password: ", "Use the password reset link on the login screen; you'll receive an email with a reset link."),
        ],
    )

    # ----- Individual Competitor -----
    add_heading(doc, "Competing as an Individual", level=1)
    add_paragraph(
        doc,
        "Your teacher will create a competition and share a short Competition Code "
        "(a six-character code such as a1b2c3). Use that code to join.",
    )
    add_heading(doc, "Join a Competition", level=2)
    add_numbered(
        doc,
        [
            "On your dashboard, choose Join Competition.",
            "Enter the Competition Code exactly as provided by your teacher.",
            "If the competition is restricted, also enter the access code your teacher gave you.",
            "Confirm. A new Competition account is created for you with $100,000 in cash.",
        ],
    )
    add_heading(doc, "What You Can See", level=2)
    add_bullets(
        doc,
        [
            "Your individual holdings (symbol, quantity, average buy price, current price, market value).",
            "Cash balance and total portfolio value.",
            "Realized P&L, unrealized P&L, and daily P&L versus start-of-day value.",
            "Return percentage, calculated as (portfolio value − $100,000) ÷ $100,000 × 100.",
            "The competition leaderboard — see where you rank against every other student in the competition.",
            "Your trade blotter (full history of executed trades for this competition).",
        ],
    )

    # ----- Team -----
    add_heading(doc, "Competing as a Team", level=1)
    add_paragraph(
        doc,
        "Teams let a group of students manage one shared portfolio. Every teammate "
        "can buy and sell from the same $100,000 pool, so coordination matters.",
    )

    add_heading(doc, "Create or Join a Team", level=2)
    add_bullets(
        doc,
        [
            ("Create a team: ", "One student creates the team and receives a Team Code (the team's ID). Share this with teammates."),
            ("Join a team: ", "Other students enter the Team Code to join. You can only join a team once."),
            ("Join a competition as a team: ", "After the team is formed, the team joins a competition using the same Competition Code an individual would use. A single Team account is created for the whole team with $100,000."),
        ],
    )

    add_heading(doc, "How Team Trading Works", level=2)
    add_bullets(
        doc,
        [
            "Any teammate can place buy or sell orders on behalf of the team.",
            "Every trade draws from — or deposits into — the shared team cash balance.",
            "All teammates see the same holdings, cash, and P&L in real time.",
            "The team has its own leaderboard, ranking team portfolios against each other.",
        ],
    )

    add_callout(
        doc,
        "Tip",
        "Agree on a trading plan with your teammates before placing orders. Because "
        "cash is shared, uncoordinated trades can drain the account or double up on "
        "positions unintentionally.",
    )

    # ----- Trading -----
    add_heading(doc, "Placing Trades", level=1)

    add_heading(doc, "Order Types", level=2)
    add_bullets(
        doc,
        [
            ("Market order: ", "Buys or sells immediately at the current live price. Fills instantly as long as you have the cash (to buy) or the shares (to sell)."),
            ("Limit order: ", "Sets a target price. A buy limit fills only at or below your limit; a sell limit fills only at or above it. Orders remain open until filled, expired, or cancelled."),
        ],
    )

    add_heading(doc, "Submitting an Order", level=2)
    add_numbered(
        doc,
        [
            "Select the account you want to trade in: Global, a Competition, or a Team.",
            "Search for a ticker symbol (e.g. AAPL, MSFT) and review the Stock Overview.",
            "Choose Buy or Sell, then enter the quantity and order type.",
            "For a limit order, enter the limit price and review it carefully.",
            "Confirm. Market orders fill right away; limit orders appear under Open Orders.",
        ],
    )

    add_heading(doc, "Trading Rules and Limits", level=2)
    add_bullets(
        doc,
        [
            ("Competition window: ", "Trades are only accepted between the competition's start and end dates. Orders placed outside the window are rejected."),
            ("Position limits: ", "A competition may cap how much of your portfolio any one stock can represent (for example, 50%). If a buy would push you over the limit, it is rejected with a message explaining why."),
            ("Cash requirement: ", "You must have enough cash to cover the full cost of a buy. There is no margin or borrowing."),
            ("No shorting or options: ", "Only long positions in stocks are supported. You cannot short sell or trade options."),
            ("Live prices: ", "Quotes come from a live market data feed, so prices move during market hours."),
        ],
    )

    add_heading(doc, "Managing Limit Orders", level=2)
    add_bullets(
        doc,
        [
            "Open Orders shows every unfilled limit order for the selected account.",
            "Orders may be open, partially filled, filled, cancelled, expired, or rejected.",
            "You can cancel an open limit order at any time before it fills.",
        ],
    )

    # ----- Portfolio -----
    add_heading(doc, "Your Portfolio Dashboard", level=1)
    add_bullets(
        doc,
        [
            "Cash balance and total portfolio value (cash + market value of holdings).",
            "Per-position detail: quantity, average buy price, current price, market value, and unrealized P&L.",
            "Realized P&L from positions you have closed and unrealized P&L from positions you still hold.",
            "Daily P&L versus the snapshot taken at the start of today.",
            "Return percentage since you started with $100,000.",
            "Trade blotter with the full history of executed trades for the selected account.",
            "Performance history with daily snapshots you can chart over time.",
        ],
    )

    # ----- Leaderboards -----
    add_heading(doc, "Leaderboards and Scoring", level=1)
    add_paragraph(
        doc,
        "Competitions are scored on total portfolio value, which is simply your cash "
        "plus the current market value of everything you hold. Because every player "
        "starts with the same $100,000, the leaderboard effectively shows who has "
        "grown their account the most.",
    )
    add_bullets(
        doc,
        [
            ("Individual leaderboard: ", "Ranks every student in the competition by total portfolio value, with their P&L and return percentage."),
            ("Team leaderboard: ", "Ranks teams in the competition by total team portfolio value."),
            ("Return %: ", "Calculated as (portfolio value − $100,000) ÷ $100,000 × 100."),
        ],
    )

    # ----- Research -----
    add_heading(doc, "Researching Stocks", level=1)
    add_bullets(
        doc,
        [
            "Search any ticker to see a Stock Overview with the current price, previous close, daily change, and day/52-week range.",
            "Toggle the chart range (1D, 1W, 1M, 6M, 1Y) to view historical price action.",
            "Use this view to evaluate a trade before placing an order from the same screen.",
        ],
    )

    # ----- Curriculum -----
    add_heading(doc, "Lessons and Assignments (Optional)", level=1)
    add_paragraph(
        doc,
        "If your teacher has attached a curriculum to the competition, you'll see "
        "lessons, readings, and assignments alongside the trading tools. Quizzes are "
        "graded automatically; written assignments are graded by your teacher.",
    )

    # ----- FAQ -----
    add_heading(doc, "Frequently Asked Questions", level=1)

    faq = [
        (
            "Do I have to join a competition to trade?",
            "No. Your Global account always lets you practice trading with $100,000, even if you have not joined a competition.",
        ),
        (
            "Can I be in more than one competition at once?",
            "Yes. Each competition creates its own separate account, with its own $100,000 starting balance.",
        ),
        (
            "Can I be on a team and compete as an individual at the same time?",
            "Yes. Your Team account and your individual Competition account are independent, so trades in one have no effect on the other.",
        ),
        (
            "What happens if I run out of cash?",
            "You simply cannot place new buys until you sell something to free up cash. There is no borrowing or margin.",
        ),
        (
            "Why was my order rejected?",
            "Common reasons: the competition window has not started or has ended, you don't have enough cash or shares, or the buy would exceed the competition's per-position limit.",
        ),
        (
            "Can I cancel a trade after it fills?",
            "No. Executed market and limit orders are final. You can only cancel an open limit order before it fills.",
        ),
    ]
    for question, answer in faq:
        q = doc.add_paragraph()
        q_run = q.add_run(question)
        q_run.bold = True
        q_run.font.color.rgb = ACCENT
        doc.add_paragraph(answer)

    # ----- Good luck -----
    add_heading(doc, "Good Luck!", level=1)
    add_paragraph(
        doc,
        "Think carefully, diversify, and manage risk. The best players combine "
        "research with discipline — not just picking winners, but knowing when to "
        "take profits and when to cut losses. Have fun, and may the best trader (or "
        "team!) win.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
